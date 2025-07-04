from flask import Flask, request, send_file, render_template, jsonify, Response, session
import json
import base64
import io
import sqlite3
import os
from datetime import datetime, timedelta
from weasyprint import HTML, CSS
from openai import OpenAI
import httpx
import uuid
import requests
import zipfile
import shutil
import csv
import logging
from docx import Document
from docx.shared import Inches, Pt
from docx.oxml.ns import qn
from PIL import Image, ImageDraw, ImageFont
import pytesseract
import urllib.parse
import re
from termcolor import colored
import werkzeug.serving  # For detecting Flask reloader
import socket  # For getting network IP address

app = Flask(__name__)
app.secret_key = os.urandom(24)  # For session management

# Set up logging
logging.basicConfig(level=logging.DEBUG, filename='app.log', filemode='a')
logger = logging.getLogger(__name__)

# Progress tracking
progress_tracker = {}

# Font and image setup
OTF_FONT_PATH = "./NotoSerifCJKsc-Regular.otf"
ZIP_URL = "https://github.com/notofonts/noto-cjk/releases/download/Serif2.003/09_NotoSerifCJKsc.zip"
ZIP_PATH = "./09_NotoSerifCJKsc.zip"
DEFAULT_IMAGE_PATH = "./static/default_image.png"

def get_network_ip():
    """Get the network-assigned IP address of the machine."""
    try:
        # Create a socket and connect to an external server to get the local IP
        s = socket.socket(socket.AF_INET, socket.SOCK_DGRAM)
        s.settimeout(0)
        s.connect(('8.8.8.8', 80))  # Use Google's DNS server as a target
        ip = s.getsockname()[0]
        s.close()
        return ip
    except Exception as e:
        logger.error("Failed to get network IP: %s", str(e))
        return "unknown"

def print_startup_message(port):
    # Only print if not running in Flask's reloader
    if werkzeug.serving.is_running_from_reloader():
        return
    ascii_art = """
╔══════════════════════════════════════╗
║                                      ║
║     #     #  # # # #   # # # #       ║
║     #     #  #     #   #             ║
║     # # # #  #     #   # # # #       ║
║     #     #  #     #         #       ║
║     #     #  # # # #   # # # #       ║
║                                      ║
╚══════════════════════════════════════╝
"""
    local_ip = "127.0.0.1"
    network_ip = get_network_ip()
    print(colored(ascii_art, 'blue'))
    print(colored("🚀 实习报告生成系统已启动！", 'cyan'))
    print(colored(f"🌐 访问地址 (本地): http://{local_ip}:{port}", 'green'))
    if network_ip != "unknown":
        print(colored(f"🌐 访问地址 (网络): http://{network_ip}:{port}", 'green'))
    else:
        print(colored("⚠️ 无法获取网络IP地址，请检查网络连接", 'yellow'))
    print(colored("📋 请复制以上链接到浏览器以访问网页界面", 'yellow'))
    print(colored("🔧 按 Ctrl+C 停止服务器", 'magenta'))
    logging.info(f"实习报告生成系统启动在 http://{local_ip}:{port} 和 http://{network_ip}:{port}")

def download_and_extract_font():
    if os.path.exists(OTF_FONT_PATH):
        logger.info("Using existing NotoSerifCJKsc-Regular.otf")
        return OTF_FONT_PATH
    
    logger.info("Downloading NotoSerifCJKsc zip file...")
    try:
        response = requests.get(ZIP_URL, stream=True)
        response.raise_for_status()
        with open(ZIP_PATH, "wb") as f:
            for chunk in response.iter_content(chunk_size=8192):
                f.write(chunk)
        logger.info("Zip file downloaded successfully.")

        with zipfile.ZipFile(ZIP_PATH, 'r') as zip_ref:
            for file_info in zip_ref.infolist():
                if file_info.filename == 'OTF/SimplifiedChinese/NotoSerifCJKsc-Regular.otf':
                    zip_ref.extract(file_info, './')
                    os.rename('./OTF/SimplifiedChinese/NotoSerifCJKsc-Regular.otf', OTF_FONT_PATH)
                    logger.info("NotoSerifCJKsc-Regular.otf extracted successfully.")
                    break
            else:
                raise Exception("NotoSerifCJKsc-Regular.otf not found in zip file.")
        
        if os.path.exists('./OTF'):
            shutil.rmtree('./OTF')
        if os.path.exists(ZIP_PATH):
            os.remove(ZIP_PATH)
        return OTF_FONT_PATH
    except Exception as e:
        logger.error(f"Failed to download or extract NotoSerifCJKsc-Regular.otf: {e}")
        return None

def create_default_image():
    if not os.path.exists(DEFAULT_IMAGE_PATH):
        os.makedirs('./static', exist_ok=True)
        img = Image.new('RGB', (200, 200), color='lightgray')
        d = ImageDraw.Draw(img)
        try:
            font = ImageFont.truetype(OTF_FONT_PATH, 20) if os.path.exists(OTF_FONT_PATH) else ImageFont.load_default()
            d.text((10, 90), "默认图片", fill='black', font=font)
        except Exception as e:
            logger.error(f"Failed to load font for default image: {e}")
            d.text((10, 90), "默认图片", fill='black')
        try:
            img.save(DEFAULT_IMAGE_PATH, format='PNG')
            logger.info("Default image created at %s", DEFAULT_IMAGE_PATH)
        except Exception as e:
            logger.error(f"Failed to save default image: {e}")
            raise

def validate_image(image_stream=None, image_path=None):
    try:
        if image_stream:
            image_stream.seek(0)
            Image.open(image_stream).verify()
            image_stream.seek(0)
            logger.debug("Uploaded image validated successfully")
            return True
        elif image_path and os.path.exists(image_path):
            Image.open(image_path).verify()
            logger.debug("Default image validated successfully at %s", image_path)
            return True
        else:
            logger.error("No valid image provided: path=%s exists=%s", image_path, os.path.exists(image_path) if image_path else False)
            return False
    except Exception as e:
        logger.error("Image validation failed: %s", str(e))
        return False

def get_image_dimensions(image_stream):
    try:
        image_stream.seek(0)
        img = Image.open(image_stream)
        width, height = img.size
        image_stream.seek(0)
        return width, height
    except Exception as e:
        logger.error("Failed to get image dimensions: %s", str(e))
        return None, None

def extract_ocr_text(image_stream):
    try:
        image_stream.seek(0)
        img = Image.open(image_stream)
        text = pytesseract.image_to_string(img, lang='chi_sim')
        image_stream.seek(0)
        logger.debug("OCR text extracted: %s", text[:100] + "..." if len(text) > 100 else text)
        return text.strip()
    except Exception as e:
        logger.error("OCR extraction failed: %s", str(e))
        return ""

def load_api_config():
    try:
        conn = sqlite3.connect(DB_PATH)
        cursor = conn.cursor()
        cursor.execute('SELECT deepseek_api_key, wecom_corpid, wecom_app_secret, wecom_agent_id, wecom_recipients FROM wecom_config WHERE id = 1')
        result = cursor.fetchone()
        conn.close()
        if result:
            config = {
                'deepseek_api_key': result[0],
                'wecom_corpid': result[1],
                'wecom_app_secret': result[2],
                'wecom_agent_id': result[3],
                'wecom_recipients': result[4]
            }
            logger.debug("API config loaded from database")
            return config
        else:
            logger.warning("No API config found in database")
            return None
    except Exception as e:
        logger.error("Failed to load API config: %s", str(e))
        return None

def save_api_config(deepseek_api_key, wecom_corpid, wecom_app_secret, wecom_agent_id, wecom_recipients):
    try:
        conn = sqlite3.connect(DB_PATH)
        cursor = conn.cursor()
        cursor.execute('''
            INSERT OR REPLACE INTO wecom_config (id, deepseek_api_key, wecom_corpid, wecom_app_secret, wecom_agent_id, wecom_recipients)
            VALUES (1, ?, ?, ?, ?, ?)
        ''', (deepseek_api_key, wecom_corpid, wecom_app_secret, wecom_agent_id, wecom_recipients))
        conn.commit()
        conn.close()
        logger.info("API config saved to database")
    except Exception as e:
        logger.error("Failed to save API config: %s", str(e))
        raise

def get_wecom_access_token(corpid, app_secret):
    try:
        response = requests.get(
            f"https://qyapi.weixin.qq.com/cgi-bin/gettoken?corpid={corpid}&corpsecret={app_secret}",
            proxies=None
        )
        response.raise_for_status()
        data = response.json()
        if data.get('errcode') == 0:
            logger.debug("WeCom access token obtained")
            return data['access_token']
        else:
            logger.error("Failed to get WeCom access token: %s", data)
            raise Exception(f"WeCom token error: {data.get('errmsg')}")
    except Exception as e:
        logger.error("WeCom token request failed: %s", str(e))
        raise

def upload_wecom_media(access_token, file_stream, filename):
    try:
        file_stream.seek(0)
        response = requests.post(
            f"https://qyapi.weixin.qq.com/cgi-bin/media/upload?access_token={access_token}&type=file",
            files={'media': (filename, file_stream, 'application/vnd.openxmlformats-officedocument.wordprocessingml.document')},
            proxies=None
        )
        response.raise_for_status()
        data = response.json()
        if data.get('errcode') == 0:
            logger.debug("WeCom media uploaded, media_id: %s", data['media_id'])
            return data['media_id']
        else:
            logger.error("Failed to upload WeCom media: %s", data)
            raise Exception(f"WeCom upload error: {data.get('errmsg')}")
    except Exception as e:
        logger.error("WeCom media upload failed: %s", str(e))
        raise

def send_wecom_file(access_token, agent_id, recipients, media_id):
    try:
        payload = {
            "touser": "|".join(recipients.split(',')),
            "msgtype": "file",
            "agentid": int(agent_id),
            "file": {"media_id": media_id},
            "enable_duplicate_check": 0
        }
        response = requests.post(
            f"https://qyapi.weixin.qq.com/cgi-bin/message/send?access_token={access_token}",
            json=payload,
            proxies=None
        )
        response.raise_for_status()
        data = response.json()
        if data.get('errcode') == 0:
            logger.info("WeCom file sent to recipients: %s", recipients)
        else:
            logger.error("Failed to send WeCom file: %s", data)
            raise Exception(f"WeCom send error: {data.get('errmsg')}")
    except Exception as e:
        logger.error("WeCom file send failed: %s", str(e))
        raise

FONT_PATH = download_and_extract_font()
create_default_image()

DB_PATH = 'reports.db'

def init_database():
    try:
        conn = sqlite3.connect(DB_PATH)
        cursor = conn.cursor()
        # Create DR_reports table with ocr_text column
        cursor.execute('''
            CREATE TABLE IF NOT EXISTS DR_reports (
                id TEXT PRIMARY KEY,
                report_date TEXT,
                intern_name TEXT,
                department TEXT,
                mentor TEXT,
                work_content TEXT,
                achievements_reflections TEXT,
                issues_improvements TEXT,
                next_day_plan TEXT,
                image_data TEXT,
                ocr_text TEXT,
                created_at TEXT,
                api_response TEXT
            )
        ''')
        # Create mentors_departments table
        cursor.execute('''
            CREATE TABLE IF NOT EXISTS mentors_departments (
                id INTEGER PRIMARY KEY AUTOINCREMENT,
                mentor TEXT,
                department TEXT
            )
        ''')
        # Create wecom_config table
        cursor.execute('''
            CREATE TABLE IF NOT EXISTS wecom_config (
                id INTEGER PRIMARY KEY,
                deepseek_api_key TEXT,
                wecom_corpid TEXT,
                wecom_app_secret TEXT,
                wecom_agent_id TEXT,
                wecom_recipients TEXT
            )
        ''')
        # Seed default mentor and department
        cursor.execute('''
            INSERT OR IGNORE INTO mentors_departments (mentor, department)
            VALUES (?, ?)
        ''', ('NAME', 'APARTMENT'))
        cursor.execute("PRAGMA table_info(DR_reports)")
        columns = [info[1] for info in cursor.fetchall()]
        if 'api_response' not in columns:
            cursor.execute('ALTER TABLE DR_reports ADD COLUMN api_response TEXT')
            logger.info("Added api_response column to DR_reports table")
        if 'ocr_text' not in columns:
            cursor.execute('ALTER TABLE DR_reports ADD COLUMN ocr_text TEXT')
            logger.info("Added ocr_text column to DR_reports table")
        conn.commit()
        conn.close()
        logger.info("Database initialized at %s", DB_PATH)
    except Exception as e:
        logger.error("Failed to initialize database: %s", str(e))
        raise

if not os.path.exists(DB_PATH):
    init_database()
else:
    init_database()

@app.route('/')
def index():
    config = load_api_config()
    api_connected = bool(config and config['deepseek_api_key'])
    try:
        conn = sqlite3.connect(DB_PATH)
        cursor = conn.cursor()
        cursor.execute('SELECT DISTINCT mentor FROM mentors_departments')
        mentors = [row[0] for row in cursor.fetchall()]
        cursor.execute('SELECT DISTINCT department FROM mentors_departments')
        departments = [row[0] for row in cursor.fetchall()]
        conn.close()
        logger.debug("Rendering index with api_connected=%s, mentors=%s, departments=%s", api_connected, mentors, departments)
        return render_template('index.html', api_connected=api_connected, mentors=mentors, departments=departments, config=config or {})
    except Exception as e:
        logger.error("Failed to fetch mentors and departments: %s", str(e))
        return render_template('index.html', api_connected=api_connected, mentors=['NAME'], departments=['APARTMENT'], config=config or {})

@app.route('/favicon.ico')
def favicon():
    logger.debug("Serving favicon.ico")
    return send_file('./static/favicon.ico', mimetype='image/x-icon')

@app.route('/.well-known/appspecific/com.chrome.devtools.json')
def chrome_devtools():
    logger.debug("Ignoring Chrome DevTools request")
    return jsonify({}), 200

@app.route('/progress/<request_id>')
def get_progress(request_id):
    progress = progress_tracker.get(request_id, {"status": "not_started", "progress": 0})
    logger.debug("Progress for %s: %s", request_id, progress)
    return jsonify(progress)

@app.route('/save_config', methods=['POST'])
def save_config():
    deepseek_api_key = request.form.get('deepseek_api_key', '')
    wecom_corpid = request.form.get('wecom_corpid', '')
    wecom_app_secret = request.form.get('wecom_app_secret', '')
    wecom_agent_id = request.form.get('wecom_agent_id', '')
    wecom_recipients = request.form.get('wecom_recipients', '')
    
    if not deepseek_api_key:
        logger.error("No DeepSeek API key provided in save_config request")
        return jsonify({"status": "error", "message": "请输入DeepSeek API密钥"})
    
    logger.debug("Testing DeepSeek API with key: %s", deepseek_api_key[:4] + "..." + deepseek_api_key[-4:])
    try:
        http_client = httpx.Client(proxies=None)
        client = OpenAI(api_key=deepseek_api_key, base_url="https://api.deepseek.com/v1", http_client=http_client)
        response = client.chat.completions.create(
            model="deepseek-chat",
            messages=[{"role": "user", "content": "Hello"}],
            stream=False
        )
        if response.choices:
            logger.info("DeepSeek API test successful")
        else:
            logger.error("DeepSeek API response missing choices")
            return jsonify({"status": "error", "message": "DeepSeek API响应异常: 无有效响应"})
    except Exception as e:
        logger.error("DeepSeek API test failed: %s", str(e))
        return jsonify({"status": "error", "message": f"DeepSeek API连接失败: {str(e)}"})

    if wecom_corpid and wecom_app_secret and wecom_agent_id and wecom_recipients:
        try:
            access_token = get_wecom_access_token(wecom_corpid, wecom_app_secret)
            logger.info("WeCom API test successful")
        except Exception as e:
            logger.error("WeCom API test failed: %s", str(e))
            return jsonify({"status": "error", "message": f"WeCom API连接失败: {str(e)}"})

    try:
        save_api_config(deepseek_api_key, wecom_corpid, wecom_app_secret, wecom_agent_id, wecom_recipients)
        logger.info("Configuration saved successfully")
        return jsonify({"status": "success", "message": "配置保存成功"})
    except Exception as e:
        logger.error("Failed to save configuration: %s", str(e))
        return jsonify({"status": "error", "message": f"配置保存失败: {str(e)}"})

@app.route('/export', methods=['GET'])
def export_reports():
    logger.debug("Exporting reports to CSV")
    conn = sqlite3.connect(DB_PATH)
    cursor = conn.cursor()
    cursor.execute('''
        SELECT id, report_date, intern_name, department, mentor, work_content,
               achievements_reflections, issues_improvements, next_day_plan, image_data, ocr_text, created_at, api_response
        FROM DR_reports
        ORDER BY intern_name, created_at
    ''')
    rows = cursor.fetchall()
    conn.close()

    output = io.StringIO()
    writer = csv.writer(output, lineterminator='\n')
    writer.writerow([
        'ID', 'Report Date', 'Intern Name', 'Department', 'Mentor', 'Work Content',
        'Achievements/Reflections', 'Issues/Improvements', 'Next Day Plan', 'Image Data', 'OCR Text', 'Created At', 'API Response'
    ])
    for row in rows:
        writer.writerow(row)
    
    output.seek(0)
    logger.info("CSV export generated")
    return Response(
        output.getvalue(),
        mimetype='text/csv',
        headers={
            "Content-Disposition": f"attachment; filename=reports_{datetime.now().strftime('%Y%m%d_%H%M%S')}.csv; filename*=UTF-8''reports_{datetime.now().strftime('%Y%m%d_%H%M%S')}.csv"
        }
    )

@app.route('/weekly_report', methods=['POST'])
def weekly_report():
    logger.debug("Received weekly report request: %s", request.form)
    config = load_api_config()
    if not config or not config['deepseek_api_key']:
        logger.error("No DeepSeek API key configured")
        return jsonify({"status": "error", "message": "缺少DeepSeek API密钥，请在设置中配置"}), 400

    week_start = request.form.get('week_start')
    intern_name = request.form.get('intern_name', '[待填]')
    if not week_start:
        logger.error("No week start date provided")
        return jsonify({"status": "error", "message": "请选择周报起始日期"}), 400

    try:
        week_start_date = datetime.strptime(week_start, '%Y-%m-%d')
        week_end_date = week_start_date + timedelta(days=6)
        week_number = week_start_date.isocalendar()[1]
        year = week_start_date.year
        date_range = f"{week_start_date.strftime('%Y年%m月%d日')} - {week_end_date.strftime('%Y年%m月%d日')}"
        filename = f"实习周报-{year}-{week_number:02d}-{intern_name}.docx"
        ascii_filename = f"weekly_report_{year}-{week_number:02d}.docx"
    except ValueError as e:
        logger.error("Invalid week start date format: %s", str(e))
        return jsonify({"status": "error", "message": "无效的日期格式"}), 400

    # Query daily reports for the week
    try:
        conn = sqlite3.connect(DB_PATH)
        cursor = conn.cursor()
        cursor.execute('''
            SELECT report_date, intern_name, department, mentor, work_content
            FROM DR_reports
            WHERE intern_name = ? AND report_date BETWEEN ? AND ?
            ORDER BY report_date
        ''', (intern_name, week_start_date.strftime('%Y年%m月%d日'), week_end_date.strftime('%Y年%m月%d日')))
        reports = cursor.fetchall()
        conn.close()
        logger.debug("Found %d daily reports for week %s", len(reports), week_start)
    except Exception as e:
        logger.error("Failed to query daily reports: %s", str(e))
        return jsonify({"status": "error", "message": f"数据库查询失败: {str(e)}"}), 500

    if not reports:
        logger.warning("No daily reports found for week %s, intern %s", week_start, intern_name)
        return jsonify({"status": "error", "message": "所选周内无日报数据"}), 404

    # Prepare data for AI summary and intensity assessment
    report_data = [
        {
            "date": report[0],
            "work_content": report[4]
        } for report in reports
    ]

    # Enhance with DeepSeek API
    try:
        http_client = httpx.Client(proxies=None)
        client = OpenAI(api_key=config['deepseek_api_key'], base_url="https://api.deepseek.com/v1", http_client=http_client)
        prompt = f"""
        You are a professional report summarization assistant. Based on the following daily reports for intern {intern_name}, generate a JSON object containing:
        - summary: A verbose summary of the week's work (500-600 characters in Chinese, excluding punctuation). Do not include the intern's name in the summary text.
        - intensity_assessments: An array of objects with date, work_content (truncated to 50 characters), and intensity (Low, Medium, High) based on work_content length (<200: Low, 200-350: Medium, >350: High) and keywords (e.g., '紧急', '复杂' increase intensity). Do not include the intern's name in work_content or intensity descriptions.
        Input:
        {json.dumps(report_data, ensure_ascii=False)}
        Output Example:
        ```json
        {{
            "summary": "本周工作总结，约500-600字，详细描述工作内容、成果和经验。",
            "intensity_assessments": [
                {{"date": "2025年07月01日", "work_content": "工作内容截断到50字...", "intensity": "Medium"}}
            ]
        }}
        ```
        """
        response = client.chat.completions.create(
            model="deepseek-chat",
            messages=[
                {"role": "system", "content": "You are a professional report summarization assistant. Do not include the intern's name in generated content."},
                {"role": "user", "content": prompt}
            ],
            response_format={"type": "json_object"},
            stream=False
        )
        if response.choices and response.choices[0].message.content:
            weekly_data = json.loads(response.choices[0].message.content)
            api_response = json.dumps(response.to_dict(), ensure_ascii=False)
            logger.debug("Weekly report API response: %s", api_response[:100] + "..." if len(api_response) > 100 else api_response)
        else:
            logger.error("Weekly report API response empty or invalid")
            raise Exception("API响应为空或无效")
    except Exception as e:
        logger.error("DeepSeek API error for weekly report: %s", str(e))
        return jsonify({"status": "error", "message": f"周报生成失败: {str(e)}"}), 500

    # Validate weekly_data
    if not isinstance(weekly_data, dict) or 'summary' not in weekly_data or 'intensity_assessments' not in weekly_data:
        logger.error("Invalid weekly_data format: %s", weekly_data)
        return jsonify({"status": "error", "message": "API返回数据格式无效"}), 500

    # Extract metadata from the first report
    department = reports[0][2]
    mentor = reports[0][3]

    # Generate DOCX
    try:
        doc = Document()
        # Set default font to Noto Serif CJK SC
        doc.styles['Normal'].font.name = 'Noto Serif CJK SC'
        doc.styles['Normal']._element.rPr.rFonts.set(qn('w:eastAsia'), 'Noto Serif CJK SC')
        doc.styles['Normal'].font.size = Pt(14.4)

        doc.add_heading('实习周报', 0).runs[0].font.size = Pt(18)
        doc.add_paragraph(f'日期范围: {date_range}').runs[0].font.size = Pt(14.4)
        doc.add_paragraph(f'实习生: {intern_name}').runs[0].font.size = Pt(14.4)
        doc.add_paragraph(f'部门: {department}').runs[0].font.size = Pt(14.4)
        doc.add_paragraph(f'指导老师: {mentor}').runs[0].font.size = Pt(14.4)
        doc.add_heading('本周工作总结:', level=2).runs[0].font.size = Pt(16)
        doc.add_paragraph(weekly_data['summary'] or '无总结').runs[0].font.size = Pt(14.4)
        doc.add_heading('工作强度评估:', level=2).runs[0].font.size = Pt(16)

        # Add intensity table
        table = doc.add_table(rows=1 + len(weekly_data['intensity_assessments']), cols=3)
        table.style = 'Table Grid'
        headers = table.rows[0].cells
        headers[0].text = '日期'
        headers[1].text = '工作内容'
        headers[2].text = '工作强度'
        for cell in headers:
            for run in cell.paragraphs[0].runs:
                run.font.size = Pt(14.4)
        for i, assessment in enumerate(weekly_data['intensity_assessments']):
            row = table.rows[i + 1].cells
            row[0].text = assessment.get('date', '未知日期')
            row[1].text = assessment.get('work_content', '无内容')
            row[2].text = assessment.get('intensity', '未知')
            for cell in row:
                for run in cell.paragraphs[0].runs:
                    run.font.size = Pt(14.4)

        buffer = io.BytesIO()
        doc.save(buffer)
        buffer.seek(0)
        logger.debug("Weekly DOCX generated: %s", filename)

        # Send to WeCom if configured
        if config and all([config.get('wecom_corpid'), config.get('wecom_app_secret'), config.get('wecom_agent_id'), config.get('wecom_recipients')]):
            try:
                access_token = get_wecom_access_token(config['wecom_corpid'], config['wecom_app_secret'])
                media_id = upload_wecom_media(access_token, buffer, filename)
                send_wecom_file(access_token, config['wecom_agent_id'], config['wecom_recipients'], media_id)
                logger.info("Weekly report sent to WeCom recipients")
            except Exception as e:
                logger.error("Failed to send weekly report to WeCom: %s", str(e))
                return jsonify({"status": "error", "message": f"WeCom发送失败: {str(e)}"}), 500

        logger.info("Returning weekly DOCX file: %s", filename)
        return Response(
            buffer.getvalue(),
            mimetype='application/vnd.openxmlformats-officedocument.wordprocessingml.document',
            headers={
                "Content-Disposition": f'attachment; filename="{ascii_filename}"; filename*=UTF-8\'\'{urllib.parse.quote(filename)}'
            }
        )
    except Exception as e:
        logger.error("Weekly DOCX generation error: %s", str(e))
        return jsonify({"status": "error", "message": f"周报生成失败: {str(e)}"}), 500

@app.route('/submit', methods=['POST'])
def submit_report():
    request_id = str(uuid.uuid4())
    progress_tracker[request_id] = {"status": "starting", "progress": 0}
    logger.debug("Received submit request: %s, request_id: %s", request.form, request_id)

    current_date = datetime.strptime("2025-07-04 10:25", "%Y-%m-%d %H:%M")
    date_str = current_date.strftime("%Y-%m-%d")
    formatted_date = current_date.strftime("%Y年%m月%d日")

    mode = request.form.get('mode', 'detailed')
    config = load_api_config()
    if not config or not config['deepseek_api_key']:
        logger.error("No DeepSeek API key configured")
        progress_tracker[request_id] = {"status": "error", "progress": 100, "message": "缺少DeepSeek API密钥，请在设置中配置"}
        return jsonify({"status": "error", "message": "缺少DeepSeek API密钥，请在设置中配置", "request_id": request_id}), 400

    intern_name = request.form.get('intern_name', '[待填]')
    filename = f"实习日常报告-{date_str}-{intern_name}.docx"
    ascii_filename = f"daily_report_{date_str}.docx"

    if mode == 'simplified':
        content = request.form.get('content', '')
        department = request.form.get('department', 'APARTMENT')
        mentor = request.form.get('mentor', 'NAME')
        if not content:
            logger.error("No content provided in simplified mode")
            progress_tracker[request_id] = {"status": "error", "progress": 100, "message": "简易模式下报告内容不能为空"}
            return jsonify({"status": "error", "message": "简易模式下报告内容不能为空", "request_id": request_id}), 400
        work_content = ''
        achievements_reflections = ''
        issues_improvements = []
        next_day_plan = []
    else:
        department = request.form.get('department', '[待填]')
        mentor = request.form.get('mentor', '[待填]')
        work_content = request.form.get('work_content', '')
        achievements_reflections = request.form.get('achievements_reflections', '')
        issues_improvements = request.form.getlist('issues_improvements')
        next_day_plan = request.form.getlist('next_day_plan')
        if not (work_content and achievements_reflections):
            logger.error("Missing required fields in detailed mode")
            progress_tracker[request_id] = {"status": "error", "progress": 100, "message": "详细模式下工作内容和收获与体会不能为空"}
            return jsonify({"status": "error", "message": "详细模式下工作内容和收获与体会不能为空", "request_id": request_id}), 400

    # Save mentor and department to database
    try:
        conn = sqlite3.connect(DB_PATH)
        cursor = conn.cursor()
        cursor.execute('''
            INSERT OR IGNORE INTO mentors_departments (mentor, department)
            VALUES (?, ?)
        ''', (mentor, department))
        conn.commit()
        conn.close()
        logger.debug("Saved mentor=%s and department=%s to database", mentor, department)
    except Exception as e:
        logger.error("Failed to save mentor and department to database: %s", str(e))
        progress_tracker[request_id] = {"status": "error", "progress": 100, "message": f"数据库保存失败: {str(e)}"}
        return jsonify({"status": "error", "message": f"数据库保存失败: {str(e)}", "request_id": request_id}), 500

    # Handle multiple image uploads and OCR
    progress_tracker[request_id] = {"status": "Processing images and OCR", "progress": 20}
    image_data_list = []
    image_streams = []
    ocr_text_list = []
    if 'image' in request.files:
        images = request.files.getlist('image')
        for image in images:
            if image.filename != '':
                try:
                    image_stream = io.BytesIO(image.read())
                    if validate_image(image_stream=image_stream):
                        image_data = base64.b64encode(image_stream.getvalue()).decode('utf-8')
                        image_data_list.append(image_data)
                        image_streams.append(image_stream)
                        ocr_text = extract_ocr_text(image_stream)
                        ocr_text_list.append(ocr_text)
                        logger.debug("Image %s validated, encoded, and OCR processed", image.filename)
                    else:
                        logger.warning("Invalid image file: %s", image.filename)
                except Exception as e:
                    logger.error("Failed to process image %s: %s", image.filename, str(e))
                    continue

    # If no valid images uploaded, use default image
    if not image_streams:
        image_data_list = []
        ocr_text_list = []
        image_streams = [open(DEFAULT_IMAGE_PATH, 'rb') if validate_image(image_path=DEFAULT_IMAGE_PATH) else None]
        if image_streams[0]:
            image_streams[0] = io.BytesIO(image_streams[0].read())
            image_data_list.append(base64.b64encode(image_streams[0].getvalue()).decode('utf-8'))
            ocr_text = extract_ocr_text(image_streams[0])
            ocr_text_list.append(ocr_text)
            logger.debug("Using default image with OCR")
        else:
            image_streams = []

    # Enhance report content using DeepSeek API
    progress_tracker[request_id] = {"status": "Enhancing content with AI", "progress": 40}
    api_response = ""
    try:
        logger.debug("Initializing OpenAI client with DeepSeek API key")
        http_client = httpx.Client(proxies=None)
        client = OpenAI(api_key=config['deepseek_api_key'], base_url="https://api.deepseek.com/v1", http_client=http_client)
        
        if mode == 'simplified':
            user_content = f"""
            Please enhance the following report input into a structured JSON format with fields: work_content, achievements_reflections, issues_improvements, next_day_plan. The content must be strictly based on the provided user input and OCR-extracted text from images, without fabricating or adding details not present in the inputs. Do not include the intern's name '{intern_name}' in the generated content (work_content, achievements_reflections, issues_improvements, next_day_plan). Expand the content to ensure the total text across all fields is 800-1000 characters (excluding punctuation), with a verbose, detailed, and professional narrative in Chinese. Distribute the content as follows: work_content (~300-400 characters), achievements_reflections (~300-400 characters), issues_improvements (~100-150 characters per item, at least one item), next_day_plan (~50-100 characters per item, at least two items). Ensure the content is relevant to the input and maintains a professional tone.
            Input:
            - User Content: {content}
            - OCR Text from Images: {json.dumps(ocr_text_list, ensure_ascii=False)}
            - Intern Name: {intern_name}
            - Department: {department}
            - Mentor: {mentor}
            - Report Date: {formatted_date}
            Output Example:
            ```json
            {{
                "work_content": "基于用户输入和OCR文本的工作内容描述，约300-400字",
                "achievements_reflections": "基于用户输入和OCR文本的收获与体会，约300-400字",
                "issues_improvements": [
                    {{"description": "问题: ... 改进计划: ...，约100-150字，基于输入"}}
                ],
                "next_day_plan": [
                    "1. 计划一，约50-100字，基于输入",
                    "2. 计划二，约50-100字，基于输入"
                ]
            }}
            ```
            """
        else:
            user_content = f"""
            Please enhance the following report input into a structured JSON format with fields: work_content, achievements_reflections, issues_improvements, next_day_plan. The content must be strictly based on the provided user input and OCR-extracted text from images, without fabricating or adding details not present in the inputs. Do not include the intern's name '{intern_name}' in the generated content (work_content, achievements_reflections, issues_improvements, next_day_plan). Ensure the content is relevant to the input and maintains a professional tone in Chinese.
            Input:
            - Work Content: {work_content}
            - Achievements and Reflections: {achievements_reflections}
            - Issues and Improvements: {', '.join(issues_improvements)}
            - Next Day Plan: {', '.join(next_day_plan)}
            - OCR Text from Images: {json.dumps(ocr_text_list, ensure_ascii=False)}
            - Intern Name: {intern_name}
            - Department: {department}
            - Mentor: {mentor}
            - Report Date: {formatted_date}
            Output Example:
            ```json
            {{
                "work_content": "基于用户输入和OCR文本的优化工作内容",
                "achievements_reflections": "基于用户输入和OCR文本的优化收获与体会",
                "issues_improvements": [
                    {{"description": "问题: ... 改进计划: ...，基于输入"}}
                ],
                "next_day_plan": [
                    "1. 计划一，基于输入",
                    "2. 计划二，基于输入"
                ]
            }}
            ```
            """
        
        logger.debug("Sending initial prompt to DeepSeek API")
        response = client.chat.completions.create(
            model="deepseek-chat",
            messages=[
                {"role": "system", "content": f"You are a professional report enhancement assistant. Generate content strictly based on provided inputs without adding fabricated details. Do not include the intern's name '{intern_name}' in generated content."},
                {"role": "user", "content": user_content}
            ],
            response_format={"type": "json_object"},
            stream=False
        )
        if response.choices and response.choices[0].message.content:
            enhanced_data = json.loads(response.choices[0].message.content)
            api_response = json.dumps(response.to_dict(), ensure_ascii=False)
            logger.debug("Initial API response: %s", api_response[:100] + "..." if len(api_response) > 100 else api_response)
        else:
            logger.error("Initial API response empty or invalid")
            raise Exception("API响应为空或无效")

        # Validate and strip intern_name from generated content
        name_pattern = re.compile(re.escape(intern_name), re.IGNORECASE)
        enhanced_data['work_content'] = name_pattern.sub('', enhanced_data.get('work_content', ''))
        enhanced_data['achievements_reflections'] = name_pattern.sub('', enhanced_data.get('achievements_reflections', ''))
        enhanced_data['issues_improvements'] = [
            {"description": name_pattern.sub('', item.get('description', ''))} 
            for item in enhanced_data.get('issues_improvements', [])
        ]
        enhanced_data['next_day_plan'] = [
            name_pattern.sub('', plan) 
            for plan in enhanced_data.get('next_day_plan', [])
        ]

        json_text = json.dumps(enhanced_data, ensure_ascii=False)
        validation_prompt = f"""
        You are a professional report validation assistant. Validate the following JSON report to ensure it is complete, correctly formatted, and professionally written in Chinese. The content must strictly adhere to the provided user input and OCR-extracted text, without adding fabricated details. Ensure the intern's name '{intern_name}' does not appear in work_content, achievements_reflections, issues_improvements, or next_day_plan. For Simplified Mode, ensure the total text across work_content, achievements_reflections, issues_improvements, and next_day_plan is 800-1000 characters (excluding punctuation) with a verbose narrative. Fix any issues and return the optimized JSON with fields: work_content, achievements_reflections, issues_improvements, next_day_plan.
        Input JSON:
        {json_text}
        Input for Reference:
        - User Content (Simplified Mode) or Fields (Detailed Mode): {content if mode == 'simplified' else json.dumps({'work_content': work_content, 'achievements_reflections': achievements_reflections, 'issues_improvements': issues_improvements, 'next_day_plan': next_day_plan}, ensure_ascii=False)}
        - OCR Text: {json.dumps(ocr_text_list, ensure_ascii=False)}
        Output Example:
        ```json
        {{
            "work_content": "优化后的工作内容，约300-400字，基于输入和OCR",
            "achievements_reflections": "优化后的收获与体会，约300-400字，基于输入和OCR",
            "issues_improvements": [
                {{"description": "问题: ... 改进计划: ...，约100-150字，基于输入和OCR"}}
            ],
            "next_day_plan": [
                "1. 计划一，约50-100字，基于输入和OCR",
                "2. 计划二，约50-100字，基于输入和OCR"
            ]
        }}
        ```
        """
        logger.debug("Sending JSON text for validation to DeepSeek API")
        progress_tracker[request_id] = {"status": "Validating AI content", "progress": 60}
        validation_response = client.chat.completions.create(
            model="deepseek-chat",
            messages=[
                {"role": "system", "content": f"You are a professional report validation assistant. Ensure content is strictly based on provided inputs without fabrication. Do not include the intern's name '{intern_name}' in generated content."},
                {"role": "user", "content": validation_prompt}
            ],
            response_format={"type": "json_object"},
            stream=False
        )
        if validation_response.choices and validation_response.choices[0].message.content:
            enhanced_data = json.loads(validation_response.choices[0].message.content)
            api_response = json.dumps(validation_response.to_dict(), ensure_ascii=False)
            logger.debug("Validated API response: %s", api_response[:100] + "..." if len(api_response) > 100 else api_response)
        else:
            logger.error("Validation API response empty or invalid")
            # Proceed with initial enhanced_data
    except Exception as e:
        logger.error("DeepSeek API error: %s", str(e))
        progress_tracker[request_id] = {"status": "error", "progress": 100, "message": f"AI内容生成失败: {str(e)}"}
        if mode == 'simplified':
            enhanced_data = {
                "work_content": content,
                "achievements_reflections": "",
                "issues_improvements": [],
                "next_day_plan": []
            }
        else:
            enhanced_data = {
                "work_content": work_content,
                "achievements_reflections": achievements_reflections,
                "issues_improvements": [{"description": issue} for issue in issues_improvements],
                "next_day_plan": [f"{i+1}. {plan}" for i, plan in enumerate(next_day_plan)]
            }
        api_response = f"Error: {str(e)}"

    # Validate enhanced_data
    if not isinstance(enhanced_data, dict):
        logger.error("Invalid enhanced_data format: %s", enhanced_data)
        progress_tracker[request_id] = {"status": "error", "progress": 100, "message": "API返回数据格式无效"}
        return jsonify({"status": "error", "message": "API返回数据格式无效", "request_id": request_id}), 500
    for field in ['work_content', 'achievements_reflections', 'issues_improvements', 'next_day_plan']:
        if field not in enhanced_data:
            enhanced_data[field] = "" if field in ['work_content', 'achievements_reflections'] else []
            logger.warning("Missing %s in enhanced_data, setting default", field)
        elif enhanced_data[field] is None:
            enhanced_data[field] = "" if field in ['work_content', 'achievements_reflections'] else []
            logger.warning("%s is None in enhanced_data, setting default", field)

    # Log character count (excluding punctuation)
    def count_characters(text):
        text = re.sub(r'[^\w\s]', '', text)  # Remove punctuation
        return len(text.replace(' ', ''))  # Remove spaces
    char_counts = {
        'work_content': count_characters(enhanced_data['work_content']),
        'achievements_reflections': count_characters(enhanced_data['achievements_reflections']),
        'issues_improvements': sum(count_characters(item.get('description', '')) for item in enhanced_data['issues_improvements']),
        'next_day_plan': sum(count_characters(plan) for plan in enhanced_data['next_day_plan'])
    }
    total_chars = sum(char_counts.values())
    logger.debug("Enhanced data for DOCX: %s", json.dumps(enhanced_data, ensure_ascii=False))
    logger.debug("Character counts (excluding punctuation): %s, Total: %d", char_counts, total_chars)

    # Structure JSON report
    report = {
        "type": "object",
        "properties": {
            "report_date": {
                "type": "string",
                "description": f"实习日常工作报告日期: {formatted_date}"
            },
            "intern_name": {
                "type": "string",
                "description": f"实习生: {intern_name}"
            },
            "department": {
                "type": "string",
                "description": f"部门: {department}"
            },
            "mentor": {
                "type": "string",
                "description": f"指导老师: {mentor}"
            },
            "work_content": {
                "type": "object",
                "properties": {
                    "description": {
                        "type": "string",
                        "description": "今日工作内容",
                        "value": enhanced_data["work_content"]
                    }
                }
            },
            "achievements_and_reflections": {
                "type": "object",
                "properties": {
                    "description": {
                        "type": "string",
                        "description": "收获与体会",
                        "value": enhanced_data["achievements_reflections"]
                    }
                }
            },
            "issues_and_improvements": {
                "type": "array",
                "items": {
                    "type": "object",
                    "properties": {
                        "description": {
                            "type": "string",
                            "description": "存在问题与改进计划",
                            "value": ""
                        }
                    }
                },
                "value": enhanced_data["issues_improvements"]
            },
            "next_day_plan": {
                "type": "array",
                "items": {
                    "type": "string",
                    "description": "明日计划"
                },
                "value": enhanced_data["next_day_plan"]
            },
            "date": {
                "type": "string",
                "description": f"日期: {formatted_date}"
            },
            "images": {
                "type": "array",
                "description": "上传的图片 (Base64)",
                "value": image_data_list
            },
            "ocr_text": {
                "type": "array",
                "description": "OCR提取的文本",
                "value": ocr_text_list
            }
        },
        "required": [
            "report_date",
            "intern_name",
            "department",
            "mentor",
            "work_content",
            "achievements_and_reflections",
            "issues_and_improvements",
            "next_day_plan",
            "date",
            "images",
            "ocr_text"
        ],
        "additionalProperties": False
    }

    # Save to database
    try:
        conn = sqlite3.connect(DB_PATH)
        cursor = conn.cursor()
        report_id = str(uuid.uuid4())
        cursor.execute('''
            INSERT INTO DR_reports (
                id, report_date, intern_name, department, mentor, work_content,
                achievements_reflections, issues_improvements, next_day_plan, image_data, ocr_text, created_at, api_response
            ) VALUES (?, ?, ?, ?, ?, ?, ?, ?, ?, ?, ?, ?, ?)
        ''', (
            report_id,
            formatted_date,
            intern_name,
            department,
            mentor,
            enhanced_data["work_content"],
            enhanced_data["achievements_reflections"],
            json.dumps(enhanced_data["issues_improvements"], ensure_ascii=False),
            json.dumps(enhanced_data["next_day_plan"], ensure_ascii=False),
            json.dumps(image_data_list, ensure_ascii=False),
            json.dumps(ocr_text_list, ensure_ascii=False),
            formatted_date,
            api_response
        ))
        conn.commit()
        conn.close()
        logger.info("Report saved to database with ID: %s", report_id)
    except Exception as e:
        logger.error("Database error: %s", str(e))
        progress_tracker[request_id] = {"status": "error", "progress": 100, "message": f"数据库保存失败: {str(e)}"}
        return jsonify({"status": "error", "message": f"数据库保存失败: {str(e)}", "request_id": request_id}), 500

    # Generate DOCX
    progress_tracker[request_id] = {"status": "Generating DOCX", "progress": 80}
    try:
        doc = Document()
        # Set default font to Noto Serif CJK SC
        doc.styles['Normal'].font.name = 'Noto Serif CJK SC'
        doc.styles['Normal']._element.rPr.rFonts.set(qn('w:eastAsia'), 'Noto Serif CJK SC')
        doc.styles['Normal'].font.size = Pt(14.4)

        doc.add_heading('实习日常工作报告', 0).runs[0].font.size = Pt(18)
        doc.add_paragraph(f'日期: {formatted_date}').runs[0].font.size = Pt(14.4)
        doc.add_paragraph(f'实习生: {intern_name}').runs[0].font.size = Pt(14.4)
        doc.add_paragraph(f'部门: {department}').runs[0].font.size = Pt(14.4)
        doc.add_paragraph(f'指导老师: {mentor}').runs[0].font.size = Pt(14.4)
        doc.add_heading('今日工作内容:', level=2).runs[0].font.size = Pt(16)
        doc.add_paragraph(enhanced_data['work_content'] or '无内容').runs[0].font.size = Pt(14.4)
        doc.add_heading('收获与体会:', level=2).runs[0].font.size = Pt(16)
        doc.add_paragraph(enhanced_data['achievements_reflections'] or '无内容').runs[0].font.size = Pt(14.4)
        doc.add_heading('存在问题与改进计划:', level=2).runs[0].font.size = Pt(16)
        for item in enhanced_data['issues_improvements'] or []:
            doc.add_paragraph(item.get('description', '无内容') or '无内容').runs[0].font.size = Pt(14.4)
        doc.add_heading('明日计划:', level=2).runs[0].font.size = Pt(16)
        for plan in enhanced_data['next_day_plan'] or []:
            doc.add_paragraph(plan or '无内容').runs[0].font.size = Pt(14.4)
        doc.add_heading('图片:', level=2).runs[0].font.size = Pt(16)
        
        image_added = False
        for idx, image_stream in enumerate(image_streams, 1):
            try:
                if image_stream and validate_image(image_stream=image_stream):
                    image_stream.seek(0)
                    width, height = get_image_dimensions(image_stream)
                    if width and height:
                        # Calculate target width
                        max_width = Inches(4.0)  # Max width ~10.16 cm
                        min_width = Inches(1.5)  # Min width ~3.81 cm
                        aspect_ratio = height / width
                        target_width = min(max_width, max(min_width, Inches(width / 100)))  # Rough pixel-to-inch conversion
                        target_height = target_width * aspect_ratio
                        doc.add_picture(image_stream, width=target_width)
                        doc.add_paragraph(f'图片 {idx}').runs[0].font.size = Pt(14.4)
                        image_added = True
                        logger.debug("Image %d added to DOCX with width=%s, height=%s", idx, target_width, target_height)
                    else:
                        logger.warning("Could not determine dimensions for image %d", idx)
                else:
                    logger.warning("Invalid image stream for image %d", idx)
            except Exception as e:
                logger.error("Image %d processing error in DOCX: %s", idx, str(e))
                doc.add_paragraph(f"无法添加图片 {idx}: {str(e)}").runs[0].font.size = Pt(14.4)

        if not image_added:
            logger.warning("No valid images available, adding placeholder text")
            doc.add_paragraph("无有效图片").runs[0].font.size = Pt(14.4)

        buffer = io.BytesIO()
        doc.save(buffer)
        buffer.seek(0)
        logger.debug("DOCX generated successfully, image_added=%s", image_added)

        # Send to WeCom if configured
        progress_tracker[request_id] = {"status": "Sending to WeCom", "progress": 90}
        if config and all([config.get('wecom_corpid'), config.get('wecom_app_secret'), config.get('wecom_agent_id'), config.get('wecom_recipients')]):
            try:
                access_token = get_wecom_access_token(config['wecom_corpid'], config['wecom_app_secret'])
                media_id = upload_wecom_media(access_token, buffer, filename)
                send_wecom_file(access_token, config['wecom_agent_id'], config['wecom_recipients'], media_id)
                logger.info("Daily report sent to WeCom recipients")
            except Exception as e:
                logger.error("Failed to send daily report to WeCom: %s", str(e))
                progress_tracker[request_id] = {"status": "error", "progress": 100, "message": f"WeCom发送失败: {str(e)}"}
                return jsonify({"status": "error", "message": f"WeCom发送失败: {str(e)}", "request_id": request_id}), 500

        progress_tracker[request_id] = {"status": "completed", "progress": 100}
        logger.info("Returning daily DOCX file: %s", filename)
        return Response(
            buffer.getvalue(),
            mimetype='application/vnd.openxmlformats-officedocument.wordprocessingml.document',
            headers={
                "Content-Disposition": f'attachment; filename="{ascii_filename}"; filename*=UTF-8\'\'{urllib.parse.quote(filename)}',
                "X-Request-ID": request_id
            }
        )
    except Exception as e:
        logger.error("DOCX generation error: %s", str(e))
        progress_tracker[request_id] = {"status": "error", "progress": 100, "message": f"DOCX生成失败: {str(e)}"}
        return jsonify({"status": "error", "message": f"DOCX生成失败: {str(e)}", "request_id": request_id}), 500

if __name__ == '__main__':
    os.environ.pop('HTTP_PROXY', None)
    os.environ.pop('HTTPS_PROXY', None)
    print_startup_message(5000)
    app.run(host='0.0.0.0', port=5000, debug=True)
