# -*- coding: utf-8 -*-
from flask import Flask, request, jsonify, render_template, send_file
from PIL import Image, ImageEnhance, ImageFilter
import easyocr
import numpy as np
import sqlite3
from datetime import datetime, timedelta
import os
import docx
from docx.shared import Inches
from openai import OpenAI
import httpx
from werkzeug.utils import secure_filename
import logging
import io
import re
import argparse
import socket
import sys
from termcolor import colored

app = Flask(__name__)

# 设置文件大小限制（16MB）
app.config['MAX_CONTENT_LENGTH'] = 16 * 1024 * 1024

# 配置日志
logging.basicConfig(level=logging.DEBUG, format='%(asctime)s - %(levelname)s - %(message)s')

# 使用绝对路径
BASE_DIR = os.path.abspath(os.path.dirname(__file__))
UPLOAD_FOLDER = os.path.join(BASE_DIR, 'Uploads')
if not os.path.exists(UPLOAD_FOLDER):
    os.makedirs(UPLOAD_FOLDER, exist_ok=True)
app.static_folder = BASE_DIR
app.add_url_rule('/Uploads/<path:filename>', endpoint='uploads_static', view_func=app.send_static_file)

# 初始化 EasyOCR
reader = easyocr.Reader(['ch_sim', 'en'], gpu=False)  # 支持中英文，禁用 GPU 以兼容更多环境

# 数据库初始化
def init_db():
    conn = sqlite3.connect(os.path.join(BASE_DIR, 'error_logs.db'))
    c = conn.cursor()
    c.execute('''CREATE TABLE IF NOT EXISTS errors (
                 id INTEGER PRIMARY KEY AUTOINCREMENT,
                 error_text TEXT,
                 solution TEXT,
                 timestamp TEXT,
                 image_path TEXT
                 )''')
    c.execute('''CREATE TABLE IF NOT EXISTS config (
                 key TEXT PRIMARY KEY,
                 value TEXT
                 )''')
    conn.commit()
    conn.close()

# 获取 DeepSeek API 密钥
def get_deepseek_api_key():
    conn = sqlite3.connect(os.path.join(BASE_DIR, 'error_logs.db'))
    c = conn.cursor()
    c.execute("SELECT value FROM config WHERE key = 'DEEPSEEK_API_KEY'")
    result = c.fetchone()
    conn.close()
    return result[0] if result else None

# 保存 DeepSeek API 密钥
def save_deepseek_api_key(api_key):
    conn = sqlite3.connect(os.path.join(BASE_DIR, 'error_logs.db'))
    c = conn.cursor()
    c.execute("INSERT OR REPLACE INTO config (key, value) VALUES ('DEEPSEEK_API_KEY', ?)", (api_key,))
    conn.commit()
    conn.close()

# 验证 DeepSeek API 连接
def verify_deepseek_api():
    api_key = get_deepseek_api_key()
    if not api_key:
        logging.error("未配置 DeepSeek API 密钥")
        return False, "未配置 DeepSeek API 密钥，请在配置页面设置"
    
    try:
        os.environ.pop('http_proxy', None)
        os.environ.pop('https_proxy', None)
        os.environ.pop('all_proxy', None)
        client = OpenAI(
            api_key=api_key,
            base_url="https://api.deepseek.com",
            http_client=httpx.Client(proxies=None)
        )
        response = client.chat.completions.create(
            model="deepseek-chat",
            messages=[
                {"role": "system", "content": "你是一位软件错误排查专家。请为给定的错误信息提供简洁的解决方案（用中文），输出格式为 Markdown，确保结构清晰，包含问题分析、解决方案和总结三个部分。"},
                {"role": "user", "content": "测试连接"}
            ],
            stream=False
        )
        logging.info("DeepSeek API 连接成功")
        return True, "API 连接成功"
    except Exception as e:
        logging.error(f"DeepSeek API 连接失败: {str(e)}")
        return False, f"API 连接失败: {str(e)}。请检查密钥或网络设置"

# 初始化数据库
init_db()

# 允许的图片扩展名
ALLOWED_EXTENSIONS = {'png', 'jpg', 'jpeg'}

def allowed_file(filename):
    return '.' in filename and filename.rsplit('.', 1)[1].lower() in ALLOWED_EXTENSIONS

# OCR 处理
def perform_ocr(image):
    try:
        # 图片预处理：灰度化、增强对比度、亮度调整、去噪
        image = image.convert('L')  # 转为灰度图
        image = image.filter(ImageFilter.GaussianBlur(radius=1))  # 高斯模糊去噪
        enhancer = ImageEnhance.Contrast(image)
        image = enhancer.enhance(2.0)  # 提高对比度
        enhancer = ImageEnhance.Brightness(image)
        image = enhancer.enhance(1.2)  # 提高亮度
        image = image.point(lambda p: p > 128 and 255)  # 自适应二值化
        
        # 转换为 EasyOCR 支持的格式（numpy 数组）
        image_np = np.array(image)
        if image_np is None:
            raise ValueError("图像转换失败，numpy 数组为空")
        
        # 使用 EasyOCR 进行识别
        result = reader.readtext(image_np, detail=0, paragraph=True)
        text = ' '.join(result)
        
        # 后处理：清理特殊字符、时间戳、IP 地址、系统状态、UI 元素
        text = re.sub(r'[^\x00-\x7F\x80-\xFFFF]', '', text)  # 移除非 ASCII 和非中文字符
        text = re.sub(r'\d{4}-\d{2}-\d{2}\s+\d{2}:\d{2}(:\d{2})?', '', text)  # 移除日期时间
        text = re.sub(r'\d{1,3}\.\d{1,3}\.\d{1,3}\.\d{1,3}(:\d+)?', '', text)  # 移除 IP 地址
        text = re.sub(r'\[\d+:\w+\]', '', text)  # 移除内存地址
        text = re.sub(r'\b(CPU|MEM|DISK)-\s*\d+\.?\d*%?', '', text)  # 移除系统状态
        text = re.sub(r'[\*\#®%]', '', text)  # 移除特定特殊字符
        text = re.sub(r'\b(admin|test|ceshi|HTML|WORD|EXCEL|PDF|XML)\b', '', text, flags=re.IGNORECASE)  # 移除 UI 元素和测试账户
        text = re.sub(r'\s+', ' ', text)  # 合并多余空格
        text = re.sub(r'\n+', '\n', text)  # 合并多余换行
        text = text.strip()
        
        logging.info("OCR 处理成功")
        return text if text else "无有效文字"
    except Exception as e:
        logging.error(f"OCR 错误: {str(e)}")
        return f"OCR 错误: {str(e)}"

# 调用 DeepSeek API 获取解决方案
def get_solution_from_deepseek(error_text):
    api_key = get_deepseek_api_key()
    if not api_key:
        logging.error("未配置 DeepSeek API 密钥")
        return "错误：未配置 DeepSeek API 密钥，请在配置页面设置"
    
    try:
        os.environ.pop('http_proxy', None)
        os.environ.pop('https_proxy', None)
        os.environ.pop('all_proxy', None)
        client = OpenAI(
            api_key=api_key,
            base_url="https://api.deepseek.com",
            http_client=httpx.Client(proxies=None)
        )
        response = client.chat.completions.create(
            model="deepseek-chat",
            messages=[
                {"role": "system", "content": "你是一位软件错误排查专家。请为给定的错误信息提供简洁的解决方案（用中文），输出格式为 Markdown，确保结构清晰，包含问题分析、解决方案和总结三个部分。"},
                {"role": "user", "content": f"错误信息: {error_text}\n请提供解决方案（Markdown 格式）。"}
            ],
            stream=False
        )
        logging.info("DeepSeek API 解决方案获取成功")
        return response.choices[0].message.content
    except Exception as e:
        logging.error(f"DeepSeek API 异常: {str(e)}")
        return f"API 异常: {str(e)}。请检查网络或API密钥"

# 保存错误和解决方案到数据库
def save_to_db(error_text, solution, image_path):
    timestamp = datetime.now().strftime('%Y-%m-%d %H:%M:%S')
    conn = sqlite3.connect(os.path.join(BASE_DIR, 'error_logs.db'))
    c = conn.cursor()
    c.execute("INSERT INTO errors (error_text, solution, timestamp, image_path) VALUES (?, ?, ?, ?)",
              (error_text, solution, timestamp, image_path))
    conn.commit()
    conn.close()
    logging.info(f"保存到数据库: {image_path}")

# 生成周报（DOCX 格式）
def generate_weekly_report():
    conn = sqlite3.connect(os.path.join(BASE_DIR, 'error_logs.db'))
    c = conn.cursor()
    end_date = datetime.now()
    start_date = end_date - timedelta(days=7)
    
    c.execute("SELECT * FROM errors WHERE timestamp >= ? AND timestamp <= ?",
              (start_date.strftime('%Y-%m-%d %H:%M:%S'), end_date.strftime('%Y-%m-%d %H:%M:%S')))
    records = c.fetchall()
    conn.close()

    doc = docx.Document()
    doc.add_heading('错误分析周报', 0)
    doc.add_paragraph(f"报告周期: {start_date.strftime('%Y-%m-%d')} 至 {end_date.strftime('%Y-%m-%d')}")
    
    for record in records:
        doc.add_heading(f"错误编号: {record[0]}", level=1)
        doc.add_paragraph(f"时间: {record[3]}")
        doc.add_paragraph(f"错误内容: {record[1]}")
        doc.add_paragraph(f"解决方案: {record[2]}")
        if record[4]:
            doc.add_paragraph("错误截图:")
            try:
                doc.add_picture(record[4], width=Inches(4))
            except:
                doc.add_paragraph("无法加载图片")
        doc.add_paragraph('-' * 50)
    
    report_path = os.path.join(BASE_DIR, 'Reports', f"weekly_report_{end_date.strftime('%Y%m%d')}.docx")
    if not os.path.exists(os.path.join(BASE_DIR, 'Reports')):
        os.makedirs(os.path.join(BASE_DIR, 'Reports'), exist_ok=True)
    doc.save(report_path)
    logging.info(f"周报生成: {report_path}")
    return report_path

# 生成 Markdown 格式的 OCR 文字文件
def generate_ocr_text_file(ocr_texts):
    text_content = "# 错误截图分析结果\n\n"
    for index, item in enumerate(ocr_texts, 1):
        text_content += f"## 图片 {index}: {item['filename']}\n"
        text_content += f"### 错误内容\n{item['error_text']}\n"
        text_content += f"### 解决方案\n{item['solution']}\n"
        text_content += f"### 总结\n- **关键问题**：{extract_key_issues(item['error_text'])}\n- **优先级**：{prioritize_issues(item['solution'])}\n"
        text_content += "---\n"
    
    text_file = io.BytesIO(text_content.encode('utf-8'))
    text_file.seek(0)
    return text_file

# 提取关键问题
def extract_key_issues(error_text):
    issues = []
    if "X-XSS-Protection" in error_text:
        issues.append("X-XSS-Protection 响应头缺失")
    if "Content-Security-Policy" in error_text:
        issues.append("Content-Security-Policy 响应头缺失")
    if "X-Frame-Options" in error_text:
        issues.append("X-Frame-Options 未配置")
    if "电子邮件地址" in error_text:
        issues.append("URL 中存在电子邮件地址")
    if "内存使用率" in error_text:
        issues.append("内存使用率过高")
    return "; ".join(issues) if issues else "未检测到关键问题"

# 确定优先级
def prioritize_issues(solution):
    if "X-XSS-Protection" in solution or "Content-Security-Policy" in solution or "X-Frame-Options" in solution:
        return "高 - 立即修复安全相关问题"
    elif "内存使用率" in solution:
        return "中 - 尽快优化系统资源"
    else:
        return "低 - 按计划处理"

# 获取本地 IP 地址
def get_local_ip():
    try:
        s = socket.socket(socket.AF_INET, socket.SOCK_DGRAM)
        s.connect(("8.8.8.8", 80))
        ip = s.getsockname()[0]
        s.close()
        return ip
    except Exception:
        return "127.0.0.1"

# 验证端口有效性
def is_valid_port(port):
    try:
        port = int(port)
        return 1024 <= port <= 65535
    except ValueError:
        return False

# 美化启动信息
def print_startup_message(ip, port):
    hos_ascii = """
╔══════════════════════════════════════╗
║                                      ║
║     #     #  # # # #   # # # #       ║
║     #     #  #     #   #             ║
║     # # # #  #     #   # # # #       ║
║     #     #  #     #         #       ║
║     #     #  # # # #   # # # #       ║
║                                      ║
╚══════════════════════════════════════╝
"""
    print(colored(hos_ascii, 'blue'))
    print(colored(f"🚀 HOS 错误截图分析系统已启动！", 'cyan'))
    print(colored(f"🌐 访问地址: http://{ip}:{port}", 'green'))
    print(colored(f"📋 请复制以上链接到浏览器以访问网页界面", 'yellow'))
    print(colored(f"🔧 按 Ctrl+C 停止服务器", 'magenta'))
    logging.info(f"服务器启动在 http://{ip}:{port}")

# Flask 路由
@app.route('/')
def index():
    return render_template('index.html')

@app.route('/config', methods=['GET', 'POST'])
def config():
    if request.method == 'POST':
        api_key = request.form.get('api_key')
        if api_key:
            save_deepseek_api_key(api_key)
            is_valid, message = verify_deepseek_api()
            if is_valid:
                return jsonify({'message': 'API 密钥保存成功且连接有效'})
            return jsonify({'error': f'API 密钥保存成功，但连接失败：{message}'}), 400
        return jsonify({'error': '请输入有效的 API 密钥'}), 400
    current_api_key = get_deepseek_api_key()
    return render_template('config.html', current_api_key=current_api_key)

@app.route('/upload', methods=['POST'])
def upload_image():
    logging.debug(f"收到上传请求: {request.files}")
    # 验证 API 连接（不阻止图片处理）
    is_valid, api_message = verify_deepseek_api()
    if not is_valid:
        logging.warning(f"API 验证失败: {api_message}")

    if 'image' not in request.files:
        logging.error("未提供图片")
        return jsonify({'error': '未提供图片，请选择至少一张图片'}), 400
    
    files = request.files.getlist('image')
    logging.debug(f"接收到 {len(files)} 张图片: {[f.filename for f in files if f.filename]}")
    if not files or all(not f.filename for f in files):
        logging.error("未选择任何有效图片")
        return jsonify({'error': '未选择任何有效图片，请选择图片'}), 400
    if len([f for f in files if f.filename]) > 10:
        logging.error("图片数量超过10张")
        return jsonify({'error': '最多上传10张图片'}), 400

    results = []
    for file in files:
        if not file or not file.filename:
            logging.error("文件无有效名称")
            results.append({
                'filename': '未知文件',
                'error_text': '文件无有效名称',
                'solution': '无法处理，请上传有效的图片文件',
                'image_path': None
            })
            continue
        if not allowed_file(file.filename):
            logging.error(f"无效文件类型: {file.filename}")
            results.append({
                'filename': file.filename,
                'error_text': '无效文件类型（仅支持 PNG, JPG）',
                'solution': '请上传 PNG 或 JPG 格式图片',
                'image_path': None
            })
            continue
        
        filename = secure_filename(file.filename)
        image_path = os.path.join(UPLOAD_FOLDER, filename)
        try:
            # 重试保存文件
            for _ in range(3):
                file.save(image_path)
                if os.path.exists(image_path):
                    break
            if not os.path.exists(image_path):
                logging.error(f"文件保存失败: {image_path}")
                results.append({
                    'filename': filename,
                    'error_text': '文件保存失败',
                    'solution': '请检查服务器存储空间或权限',
                    'image_path': None
                })
                continue
            
            logging.info(f"文件保存成功: {image_path}")
            
            # 执行 OCR
            image = Image.open(image_path)
            error_text = perform_ocr(image)
            
            # 获取解决方案（即使 API 失败也继续）
            solution = get_solution_from_deepseek(error_text) if is_valid else "API 连接失败，无法获取解决方案"
            
            # 保存到数据库
            save_to_db(error_text, solution, image_path)
            
            # 返回相对路径
            relative_image_path = os.path.join('Uploads', filename).replace('\\', '/')
            results.append({
                'filename': filename,
                'error_text': error_text,
                'solution': solution,
                'image_path': relative_image_path
            })
        except Exception as e:
            logging.error(f"处理文件 {filename} 失败: {str(e)}")
            results.append({
                'filename': filename,
                'error_text': f"文件处理错误: {str(e)}",
                'solution': '请检查文件格式或服务器状态',
                'image_path': None
            })
    
    logging.debug(f"上传处理完成，返回 {len(results)} 个结果")
    return jsonify({'results': results, 'api_status': api_message})

@app.route('/download_ocr_text', methods=['POST'])
def download_ocr_text():
    try:
        data = request.get_json()
        ocr_texts = data.get('results', [])
        if not ocr_texts:
            return jsonify({'error': '无 OCR 文字可下载'}), 400
        
        text_file = generate_ocr_text_file(ocr_texts)
        timestamp = datetime.now().strftime('%Y%m%d_%H%M%S')
        return send_file(
            text_file,
            mimetype='text/markdown',
            as_attachment=True,
            download_name=f'ocr_texts_{timestamp}.md'
        )
    except Exception as e:
        logging.error(f"生成 OCR 文字文件失败: {str(e)}")
        return jsonify({'error': f"生成 OCR 文字文件失败: {str(e)}"}), 500

@app.route('/generate_report', methods=['GET'])
def generate_report():
    try:
        report_path = generate_weekly_report()
        return jsonify({'report_path': report_path})
    except Exception as e:
        logging.error(f"生成周报失败: {str(e)}")
        return jsonify({'error': str(e)}), 500

if __name__ == '__main__':
    # 解析命令行参数
    parser = argparse.ArgumentParser(description='HOS 错误截图分析系统')
    parser.add_argument('--port', type=int, default=5000, help='运行端口 (1024-65535)')
    args = parser.parse_args()

    # 验证端口
    port = args.port
    if not is_valid_port(port):
        print(colored(f"错误: 无效端口 {port}，端口必须在 1024-65535 之间", 'red'))
        sys.exit(1)

    # 获取本地 IP
    local_ip = get_local_ip()

    # 打印启动信息
    print_startup_message(local_ip, port)

    # 启动 Flask 应用
    try:
        app.run(host='0.0.0.0', port=port, debug=True, use_reloader=False)
    except Exception as e:
        print(colored(f"错误: 无法启动服务器，端口 {port} 可能被占用或权限不足", 'red'))
        print(colored(f"详情: {str(e)}", 'red'))
        sys.exit(1)
